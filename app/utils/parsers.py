import os
import olefile
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
import requests
from bs4 import BeautifulSoup

def parse_hwp(file_path: str) -> str:
    try:
        if not olefile.isOleFile(file_path): return ""
        ole = olefile.OleFileIO(file_path)
        text_parts = []
        for stream in ole.listdir():
            if '/'.join(stream).startswith('BodyText/Section'):
                data = ole.openstream(stream).read()
                text = data.decode('utf-16le', errors='ignore').replace('\x00', '').strip()
                if text: text_parts.append(text)
        ole.close()
        return '\n'.join(text_parts)
    except Exception: return ""

def parse_pdf(file_path: str) -> str:
    try:
        doc = fitz.open(file_path)
        text = '\n\n'.join([page.get_text() for page in doc])
        doc.close()
        return text
    except Exception: return ""

def parse_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(' | '.join(cell.text.strip() for cell in row.cells))
        return '\n'.join(parts)
    except Exception: return ""

def parse_excel(file_path: str) -> str:
    try:
        excel_file = pd.ExcelFile(file_path)
        parts = []
        for sheet in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet)
            if not df.empty:
                parts.append(f"## 시트: {sheet}\n{df.to_markdown(index=False)}")
        return '\n'.join(parts)
    except Exception: return ""

def parse_csv(file_path: str) -> str:
    try:
        for enc in ['utf-8', 'cp949', 'euc-kr']:
            try:
                df = pd.read_csv(file_path, encoding=enc)
                return df.to_markdown(index=False) or ""
            except: continue
        return ""
    except Exception: return ""

def parse_text(file_path: str) -> str:
    try:
        for enc in ['utf-8', 'cp949', 'euc-kr']:
            try:
                with open(file_path, 'r', encoding=enc) as f: return f.read()
            except: continue
        return ""
    except Exception: return ""

def parse_web_content(url: str) -> str:
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove unwanted tags
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Extract text
        text = soup.get_text(separator='\n\n')
        
        # Clean text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        return '\n'.join(chunk for chunk in chunks if chunk)
    except Exception: return ""

def parse_file(file_path: str) -> str:
    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1].lower()
    
    if file_ext == '.hwp':
        return parse_hwp(file_path)
    elif file_ext == '.pdf':
        return parse_pdf(file_path)
    elif file_ext == '.docx':
        return parse_docx(file_path)
    elif file_ext in ['.xlsx', '.xls']:
        return parse_excel(file_path)
    elif file_ext == '.csv':
        return parse_csv(file_path)
    elif file_ext in ['.txt', '.md']:
        return parse_text(file_path)
    else:
        return ""
